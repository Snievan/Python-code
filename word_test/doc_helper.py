import docx
from typing import List, Tuple


def modify_content(doc, replace_pairs: List[Tuple]) -> None:
    """modify content by replcaing keyword"""
    all_paras = doc.paragraphs
    for i, para in enumerate(all_paras):
        text_cur = para.text
        text_new = replace_str(text_cur, replace_pairs)
        doc.paragraphs[i].text = text_new
    return doc


def modify_header(doc, replace_pairs: List[Tuple]) -> None:
    """modify header from 2nd page"""
    section = doc.sections[1]
    header = section.header
    text_cur = header.paragraphs[0].text
    text_new = replace_str(text_cur, replace_pairs)
    header.paragraphs[0].text = text_new
    return doc


def modify_table(doc, replace_pairs: List[Tuple]) -> None:
    """modify table objects"""
    indexs = [(0, 1, 0, 0), (0, 1, 0, 2)]
    for i, j, k, m in indexs:
        text_cur = doc.tables[i].cell(j, k).paragraphs[m].text
        text_new = replace_str(text_cur, replace_pairs)
        doc.tables[i].cell(j, k).paragraphs[m].text = text_new
    return doc


def gen_modified_doc(path_read: str, path_write: str, replace_pairs: List[Tuple]) -> None:
    """modify doc file and write to specific path"""
    doc = docx.Document(path_read)
    doc = modify_content(doc, replace_pairs)
    doc = modify_header(doc, replace_pairs)
    # doc = modify_table(doc, replace_pairs)
    doc.save(path_write)
    print(f'File saved at {path_write}')


def replace_str(s: str, replace_pairs: List[Tuple]) -> None:
    """replace given string by paris sequencely"""
    replace_pairs = replace_pairs.copy()
    if not replace_pairs:
        return s
    val_cur, var_new = replace_pairs.pop(0)
    s = s.replace(val_cur, var_new)
    return replace_str(s, replace_pairs=replace_pairs)


def main():
    path_read = './docs/ZB_M0027_开门输入信号异常诊断模型部署说明.docx'
    path_write = './docs/ZB_M0028_关门输入信号异常诊断模型部署说明.docx'
    replace_pairs = [('ZB_M0027', 'ZB_M0028'),
                     ('开门输入信号异常诊断模型', '关门输入信号异常诊断模型')]
    gen_modified_doc(path_read, path_write, replace_pairs)


if __name__ == "__main__":
    main()
    # docx.Document(path_read).
